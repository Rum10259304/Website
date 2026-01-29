# # ingest.py

import os
import re
from docx import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import FAISS
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain.docstore.document import Document as LangchainDocument
from config import EMBEDDING_MODEL_NAME, CHUNK_SIZE, CHUNK_OVERLAP, CLEANED_DIR, PDF_DIR, VECTORSTORE_DIR

# 🔍 Read DOCX files
def read_docx(file_path):
    doc = Document(file_path)
    return "\n".join([para.text for para in doc.paragraphs])

# 🧹 Optional: additional cleaning
def clean_text(text):
    text = re.sub(r"\*|\d+\.", "", text)
    return text.strip()

# 🚀 Main ingestion function
def ingest_documents():
    splitter = RecursiveCharacterTextSplitter(chunk_size=CHUNK_SIZE, chunk_overlap=CHUNK_OVERLAP)
    embeddings = HuggingFaceEmbeddings(model_name=EMBEDDING_MODEL_NAME)
    docs = []

    for file in os.listdir(CLEANED_DIR):
        path = os.path.join(CLEANED_DIR, file)

        # 🔄 Load text content
        if file.endswith(".txt"):
            with open(path, encoding="utf-8") as f:
                text = f.read()
            base_name = file.replace(".txt", "")
        elif file.endswith(".docx"):
            text = read_docx(path)
            base_name = file.replace(".docx", "")
        else:
            continue

        text = clean_text(text)
        chunks = splitter.create_documents([text])

        # 🔗 Match with original file
        matched_file = None
        for f in os.listdir(PDF_DIR):
            filename_wo_ext, _ = os.path.splitext(f)
            if base_name.lower() == filename_wo_ext.lower():
                matched_file = f
                break

        source_file = matched_file if matched_file else f"{base_name}.docx"

        # 🏷️ Metadata tagging
        if "cover" in base_name.lower():
            doc_type = "cover_page"
        elif any(term in base_name.lower() for term in ["digital meeting", "online meeting", "virtual meeting"]):
            doc_type = "digital"
        elif "etiquette" in base_name.lower() or "physical" in base_name.lower():
            doc_type = "physical"
        else:
            doc_type = "general"

        # 📎 Attach metadata to each chunk
        for chunk in chunks:
            chunk.metadata["source"] = source_file
            chunk.metadata["title"] = base_name.replace("_", " ").lower().strip()
            chunk.metadata["doc_type"] = doc_type
            docs.append(chunk)

    # 💾 Save to FAISS
    vectorstore = FAISS.from_documents(docs, embeddings)
    vectorstore.save_local(VECTORSTORE_DIR)
    print("✅ Ingestion complete.")

if __name__ == "__main__":
    ingest_documents()


# import os
# import re
# from docx import Document
# from langchain.text_splitter import RecursiveCharacterTextSplitter
# from langchain_community.vectorstores import FAISS
# # from langchain.embeddings import HuggingFaceEmbeddings
# from langchain_community.embeddings import HuggingFaceEmbeddings
# from langchain.docstore.document import Document as LangchainDocument
# from config import EMBEDDING_MODEL_NAME, CHUNK_SIZE, CHUNK_OVERLAP, CLEANED_DIR, PDF_DIR, VECTORSTORE_DIR

# def read_docx(file_path):
#     doc = Document(file_path)
#     return "\n".join([para.text for para in doc.paragraphs])

# def clean_text(text):
#     text = re.sub(r"\*|\d+\.", "", text)
#     return text.strip()

# def ingest_documents():
#     splitter = RecursiveCharacterTextSplitter(chunk_size=CHUNK_SIZE, chunk_overlap=CHUNK_OVERLAP)
#     embeddings = HuggingFaceEmbeddings(model_name=EMBEDDING_MODEL_NAME)
#     docs = []

#     for file in os.listdir(CLEANED_DIR):
#         path = os.path.join(CLEANED_DIR, file)

#         if file.endswith(".txt"):
#             with open(path, encoding="utf-8") as f:
#                 text = f.read()
#             base_name = file.replace(".txt", "")
#         elif file.endswith(".docx"):
#             text = read_docx(path)
#             base_name = file.replace(".docx", "")
#         else:
#             continue

#         text = clean_text(text)
#         chunks = splitter.create_documents([text])

#         # Find a matching PDF or DOCX in the PDF_DIR folder
#         # matched_file = None
#         # for f in os.listdir(PDF_DIR):
#         #     if base_name in f:
#         #         matched_file = f
#         #         break

#         matched_file = None
#         for f in os.listdir(PDF_DIR):
#             filename_wo_ext, _ = os.path.splitext(f)
#             if base_name.lower() == filename_wo_ext.lower():
#                 matched_file = f
#                 break


#         source_file = matched_file if matched_file else f"{base_name}.docx"

#         for chunk in chunks:
#             chunk.metadata["source"] = source_file
#             chunk.metadata["title"] = base_name.replace("_", " ").lower().strip()
#             docs.append(chunk)

#     vectorstore = FAISS.from_documents(docs, embeddings)
#     vectorstore.save_local(VECTORSTORE_DIR)
#     print("✅ Ingestion complete.")

# if __name__ == "__main__":
#     ingest_documents()


